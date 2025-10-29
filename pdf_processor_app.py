import streamlit as st
import fitz  # PyMuPDF
from PIL import Image, ImageDraw, ImageFont
import numpy as np
import img2pdf
import io
import zipfile
from pathlib import Path
import tempfile
import os
import time

# Try to import python-docx, but make it optional
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

st.set_page_config(page_title="PDF Image Processor 1.1", layout="wide")

st.title("🔄 PDF Image Processor 1.1")
st.markdown("**Visual Logo Selection + Both-Direction Cropping**")

if not DOCX_AVAILABLE:
    st.sidebar.warning("⚠️ Word export requires: `pip install python-docx`")

def get_all_page_images(pdf_file):
    """Extract all pages as images for logo setup"""
    pdf_data = pdf_file.getvalue()
    doc = fitz.open(stream=pdf_data, filetype="pdf")
    page_images = []
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        pix = page.get_pixmap()
        img_data = pix.tobytes("ppm")
        pil_image = Image.open(io.BytesIO(img_data))
        page_images.append(pil_image)
    
    doc.close()
    return page_images

def create_grid_overlay(image, grid_size=50):
    """Create a visible grid overlay image with coordinates"""
    # Create a semi-transparent overlay
    overlay = Image.new('RGBA', image.size, (255, 255, 255, 0))
    draw = ImageDraw.Draw(overlay)
    
    # Colors for better visibility
    grid_color = (0, 100, 255, 180)  # Blue with good opacity
    text_color = (0, 0, 0, 255)      # Solid black for text
    center_color = (255, 0, 0, 220)  # Solid red for center lines
    
    # Draw vertical lines
    for x in range(0, image.width, grid_size):
        draw.line([(x, 0), (x, image.height)], fill=grid_color, width=2)
        # Add coordinate text at top (with background for readability)
        text = str(x)
        bbox = draw.textbbox((0, 0), text)
        text_width = bbox[2] - bbox[0]
        draw.rectangle([x, 0, x + text_width + 4, 15], fill=(255, 255, 255, 200))
        draw.text((x + 2, 2), text, fill=text_color)
    
    # Draw horizontal lines
    for y in range(0, image.height, grid_size):
        draw.line([(0, y), (image.width, y)], fill=grid_color, width=2)
        # Add coordinate text at left (with background for readability)
        text = str(y)
        bbox = draw.textbbox((0, 0), text)
        text_width = bbox[2] - bbox[0]
        draw.rectangle([0, y, text_width + 4, y + 15], fill=(255, 255, 255, 200))
        draw.text((2, y + 2), text, fill=text_color)
    
    # Draw prominent center lines
    center_x = image.width // 2
    center_y = image.height // 2
    draw.line([(center_x, 0), (center_x, image.height)], fill=center_color, width=3)
    draw.line([(0, center_y), (image.width, center_y)], fill=center_color, width=3)
    
    # Add center coordinates with background
    center_text = f"Center: ({center_x}, {center_y})"
    bbox = draw.textbbox((0, 0), center_text)
    text_width = bbox[2] - bbox[0]
    draw.rectangle([center_x + 5, center_y + 5, center_x + text_width + 10, center_y + 20], 
                  fill=(255, 255, 255, 230))
    draw.text((center_x + 7, center_y + 7), center_text, fill=(255, 0, 0, 255))
    
    return overlay

def visual_logo_selection(image, logo_states):
    """Visual logo selection with interactive coordinate selection"""
    st.subheader("🎯 Logo Area Setup (4 Boxes)")
    
    # Page selection for visual reference
    if 'all_page_images' in st.session_state:
        page_options = list(range(1, len(st.session_state.all_page_images) + 1))
        selected_page = st.selectbox(
            "Select Page for Visual Reference:",
            page_options,
            index=0,
            help="Select which page to use for setting logo coordinates. Coordinates will apply to ALL pages."
        )
        
        # Update image to selected page
        image = st.session_state.all_page_images[selected_page - 1]
        st.info(f"📄 Using Page {selected_page} for reference. Logo coordinates will apply to all {len(st.session_state.all_page_images)} pages.")
    
    # Grid settings
    st.subheader("🗺️ Grid Settings")
    grid_col1, grid_col2, grid_col3 = st.columns(3)
    with grid_col1:
        show_grid = st.toggle("Show Grid Overlay", value=True, help="Display grid lines for easier coordinate selection")
    with grid_col2:
        grid_size = st.slider("Grid Size (pixels)", min_value=25, max_value=200, value=50, 
                             help="Distance between grid lines")
    with grid_col3:
        grid_opacity = st.slider("Grid Opacity", min_value=50, max_value=100, value=80,
                                help="Grid visibility level") / 100.0
    
    # Create display image with optional grid
    display_img = image.copy().convert('RGBA')
    if show_grid:
        grid_overlay = create_grid_overlay(image, grid_size)
        # Adjust opacity
        if grid_opacity < 1.0:
            grid_overlay = grid_overlay.convert('RGBA')
            data = np.array(grid_overlay)
            data[..., 3] = (data[..., 3] * grid_opacity).astype(np.uint8)
            grid_overlay = Image.fromarray(data)
        
        display_img = Image.alpha_composite(display_img, grid_overlay)
    
    # Display the reference image
    st.image(display_img, caption="Reference Image with Grid - Click buttons below to set logo areas", use_column_width=True)
    
    # Interactive Point Selection
    st.subheader("🎯 Interactive Area Selection")
    st.info("**Click the buttons below to set logo areas at common positions**")
    
    point_cols = st.columns(4)
    click_points = {}
    
    with point_cols[0]:
        if st.button("📍 Top-Left Area", use_container_width=True):
            click_points = {"x1": 10, "y1": 10, "x2": 110, "y2": 60}
    with point_cols[1]:
        if st.button("📍 Top-Right Area", use_container_width=True):
            click_points = {"x1": image.width-120, "y1": 10, "x2": image.width-20, "y2": 60}
    with point_cols[2]:
        if st.button("📍 Bottom-Left Area", use_container_width=True):
            click_points = {"x1": 10, "y1": image.height-70, "x2": 110, "y2": image.height-20}
    with point_cols[3]:
        if st.button("📍 Bottom-Right Area", use_container_width=True):
            click_points = {"x1": image.width-120, "y1": image.height-70, "x2": image.width-20, "y2": image.height-20}
    
    # Apply clicked points to Logo 1
    if click_points and not st.session_state.get('logo1_coords'):
        st.session_state.logo1_x1 = click_points["x1"]
        st.session_state.logo1_y1 = click_points["y1"]
        st.session_state.logo1_x2 = click_points["x2"]
        st.session_state.logo1_y2 = click_points["y2"]
        st.session_state.logo1_enabled = True
        st.success("✅ Logo 1 area set! Adjust coordinates below or set other logos.")
    
    # Coordinate Boxes Section
    st.subheader("📐 Coordinate Controls")
    
    # Create 4 columns for logo toggles
    toggle_cols = st.columns(4)
    logo_enabled = {}
    
    for i in range(1, 5):
        with toggle_cols[i-1]:
            logo_enabled[i] = st.toggle(f"Logo {i}", 
                                      value=logo_states[f'logo{i}_enabled'],
                                      key=f"logo{i}_toggle")
    
    # Logo coordinate inputs
    logo_coords = {}
    colors = ["red", "blue", "green", "orange"]
    
    for i in range(1, 5):
        if logo_enabled[i]:
            with st.expander(f"🎯 Logo {i} Coordinates ({colors[i-1].title()})", expanded=True):
                cols = st.columns(4)
                
                with cols[0]:
                    x1 = st.number_input(f"Left (X1)", 
                                       min_value=0, max_value=image.width,
                                       value=logo_states[f'logo{i}_coords'][0] if logo_states[f'logo{i}_coords'] else 50 + (i-1)*30,
                                       key=f"logo{i}_x1")
                with cols[1]:
                    y1 = st.number_input(f"Top (Y1)", 
                                       min_value=0, max_value=image.height,
                                       value=logo_states[f'logo{i}_coords'][1] if logo_states[f'logo{i}_coords'] else 50 + (i-1)*40,
                                       key=f"logo{i}_y1")
                with cols[2]:
                    x2 = st.number_input(f"Right (X2)", 
                                       min_value=0, max_value=image.width,
                                       value=logo_states[f'logo{i}_coords'][2] if logo_states[f'logo{i}_coords'] else 150 + (i-1)*30,
                                       key=f"logo{i}_x2")
                with cols[3]:
                    y2 = st.number_input(f"Bottom (Y2)", 
                                       min_value=0, max_value=image.height,
                                       value=logo_states[f'logo{i}_coords'][3] if logo_states[f'logo{i}_coords'] else 100 + (i-1)*40,
                                       key=f"logo{i}_y2")
                
                # Ensure valid coordinates
                if x2 <= x1:
                    x2 = x1 + 10
                if y2 <= y1:
                    y2 = y1 + 10
                
                logo_coords[i] = (x1, y1, x2, y2)
                
                # Show coordinate summary
                st.success(f"Logo {i}: ({x1}, {y1}) to ({x2}, {y2}) | Size: {x2-x1}×{y2-y1} pixels")
    
    # Real-time Preview Section
    st.subheader("👁️ Live Preview")
    
    # Create preview image with grid
    preview_img = image.copy().convert('RGBA')
    if show_grid:
        grid_overlay = create_grid_overlay(image, grid_size)
        # Adjust opacity
        if grid_opacity < 1.0:
            grid_overlay = grid_overlay.convert('RGBA')
            data = np.array(grid_overlay)
            data[..., 3] = (data[..., 3] * grid_opacity).astype(np.uint8)
            grid_overlay = Image.fromarray(data)
        
        preview_img = Image.alpha_composite(preview_img, grid_overlay)
    
    draw = ImageDraw.Draw(preview_img)
    
    active_logos = []
    for i in range(1, 5):
        if logo_enabled[i] and logo_coords.get(i):
            x1, y1, x2, y2 = logo_coords[i]
            # Draw logo box with thick border
            draw.rectangle([x1, y1, x2, y2], outline=colors[i-1], width=4)
            # Add label with background
            label = f"LOGO {i}"
            bbox = draw.textbbox((0, 0), label)
            text_width = bbox[2] - bbox[0]
            draw.rectangle([x1, y1-30, x1 + text_width + 6, y1-10], fill=(255, 255, 255, 200))
            draw.text((x1+3, y1-28), label, fill=colors[i-1])
            # Add size info with background
            size_text = f"{x2-x1}×{y2-y1}"
            bbox = draw.textbbox((0, 0), size_text)
            text_width = bbox[2] - bbox[0]
            draw.rectangle([x1, y2+2, x1 + text_width + 6, y2+22], fill=(255, 255, 255, 200))
            draw.text((x1+3, y2+4), size_text, fill=colors[i-1])
            # Add coordinates with background
            coord_text = f"({x1},{y1})"
            bbox = draw.textbbox((0, 0), coord_text)
            text_width = bbox[2] - bbox[0]
            draw.rectangle([x1, y1-60, x1 + text_width + 6, y1-40], fill=(255, 255, 255, 200))
            draw.text((x1+3, y1-58), coord_text, fill=colors[i-1])
            active_logos.append(f"Logo {i}")
    
    st.image(preview_img, caption="🔴 Logo 1 | 🔵 Logo 2 | 🟢 Logo 3 | 🟠 Logo 4", use_column_width=True)
    
    # Quick Actions
    if any(logo_enabled.values()):
        st.subheader("⚡ Quick Actions")
        action_cols = st.columns(3)
        
        with action_cols[0]:
            if st.button("📋 Copy Logo 1 to Others", use_container_width=True):
                if logo_enabled[1] and logo_coords.get(1):
                    base_x1, base_y1, base_x2, base_y2 = logo_coords[1]
                    for i in range(2, 5):
                        if logo_enabled[i]:
                            offset = (i-1) * 20
                            st.session_state[f'logo{i}_x1'] = base_x1 + offset
                            st.session_state[f'logo{i}_y1'] = base_y1 + offset
                            st.session_state[f'logo{i}_x2'] = base_x2 + offset
                            st.session_state[f'logo{i}_y2'] = base_y2 + offset
                    st.rerun()
        
        with action_cols[1]:
            if st.button("🔄 Reset All Positions", use_container_width=True):
                for i in range(1, 5):
                    if logo_enabled[i]:
                        st.session_state[f'logo{i}_x1'] = 50 + (i-1)*30
                        st.session_state[f'logo{i}_y1'] = 50 + (i-1)*40
                        st.session_state[f'logo{i}_x2'] = 150 + (i-1)*30
                        st.session_state[f'logo{i}_y2'] = 100 + (i-1)*40
                st.rerun()
        
        with action_cols[2]:
            if st.button("🎯 Auto-Space Logos", use_container_width=True):
                img_w, img_h = image.width, image.height
                logo_width, logo_height = 100, 50
                spacing_x = (img_w - (4 * logo_width)) // 5
                spacing_y = (img_h - logo_height) // 2
                
                for i in range(1, 5):
                    if logo_enabled[i]:
                        x1 = spacing_x + (i-1) * (logo_width + spacing_x)
                        y1 = spacing_y
                        st.session_state[f'logo{i}_x1'] = x1
                        st.session_state[f'logo{i}_y1'] = y1
                        st.session_state[f'logo{i}_x2'] = x1 + logo_width
                        st.session_state[f'logo{i}_y2'] = y1 + logo_height
                st.rerun()
    
    # Status and confirmation
    if active_logos:
        st.success(f"✅ Active logos: {', '.join(active_logos)}")
        st.info("💡 **Tip**: Use the grid overlay and click buttons for quick setup, then fine-tune with coordinate boxes")
    else:
        st.warning("⚠️ No logos enabled - no logo removal will be performed")
    
    if st.button("✅ Confirm All Logo Areas", type="primary"):
        # Update session state
        for i in range(1, 5):
            st.session_state[f'logo{i}_enabled'] = logo_enabled[i]
            if logo_enabled[i]:
                st.session_state[f'logo{i}_coords'] = logo_coords[i]
            else:
                st.session_state[f'logo{i}_coords'] = None
        
        st.balloons()
        st.success("🎉 All logo areas confirmed! Ready to process PDF.")
        return True
    
    return False

# ... (rest of the functions remain the same - remove_logo_precise, crop_vertical_only, crop_horizontal_only, crop_both_fixed, process_pdf_with_logos, create_word_document_with_options)

def remove_logo_precise(image, logo_coords, method="white"):
    """Remove logo with precise coordinates"""
    if logo_coords is None:
        return image
    
    x1, y1, x2, y2 = logo_coords
    
    # Create copy
    result_img = image.copy()
    draw = ImageDraw.Draw(result_img)
    
    if method == "white":
        # Simple white fill
        draw.rectangle([x1, y1, x2, y2], fill=(255, 255, 255))
    else:
        # Sample background around logo
        img_array = np.array(image)
        
        # Get samples from all sides
        samples = []
        
        # Top sample
        if y1 > 10:
            top_sample = img_array[max(0, y1-10):y1, x1:x2]
            if top_sample.size > 0:
                samples.append(np.mean(top_sample, axis=(0, 1)))
        
        # Bottom sample  
        if y2 < image.height - 10:
            bottom_sample = img_array[y2:min(image.height, y2+10), x1:x2]
            if bottom_sample.size > 0:
                samples.append(np.mean(bottom_sample, axis=(0, 1)))
        
        # Left sample
        if x1 > 10:
            left_sample = img_array[y1:y2, max(0, x1-10):x1]
            if left_sample.size > 0:
                samples.append(np.mean(left_sample, axis=(0, 1)))
        
        # Right sample
        if x2 < image.width - 10:
            right_sample = img_array[y1:y2, x2:min(image.width, x2+10)]
            if right_sample.size > 0:
                samples.append(np.mean(right_sample, axis=(0, 1)))
        
        if samples:
            fill_color = tuple(np.mean(samples, axis=0).astype(int))
        else:
            fill_color = (255, 255, 255)  # Fallback to white
        
        draw.rectangle([x1, y1, x2, y2], fill=fill_color)
    
    return result_img

def crop_vertical_only(image, white_threshold=245):
    """Crop only top and bottom white borders - INDIVIDUAL per image"""
    img_array = np.array(image)
    
    if len(img_array.shape) == 3:
        non_white_mask = np.any(img_array < white_threshold, axis=(1, 2))
    else:
        non_white_mask = np.any(img_array < white_threshold, axis=1)
    
    non_white_rows = np.where(non_white_mask)[0]
    if len(non_white_rows) == 0:
        return image
    
    y_min = non_white_rows[0]
    y_max = non_white_rows[-1]
    
    margin = 5
    y_min = max(0, y_min - margin)
    y_max = min(image.height, y_max + 1 + margin)
    
    return image.crop((0, y_min, image.width, y_max))

def crop_horizontal_only(image, white_threshold=245):
    """Crop only left and right white borders - INDIVIDUAL per image"""
    img_array = np.array(image)
    
    if len(img_array.shape) == 3:
        non_white_mask = np.any(img_array < white_threshold, axis=(0, 1))
    else:
        non_white_mask = np.any(img_array < white_threshold, axis=0)
    
    non_white_cols = np.where(non_white_mask)[0]
    if len(non_white_cols) == 0:
        return image
    
    x_min = non_white_cols[0]
    x_max = non_white_cols[-1]
    
    margin = 5
    x_min = max(0, x_min - margin)
    x_max = min(image.width, x_max + 1 + margin)
    
    return image.crop((x_min, 0, x_max, image.height))

def crop_both_fixed(image, white_threshold=245):
    """FIXED: Crop both vertical and horizontal - INDIVIDUAL per image"""
    img_array = np.array(image)
    
    # Create a unified mask of non-white areas for THIS specific image
    if len(img_array.shape) == 3:
        # For RGB images: pixel is white if ALL channels are above threshold
        white_mask = np.all(img_array >= white_threshold, axis=2)
    else:
        # For grayscale images
        white_mask = img_array >= white_threshold
    
    # Find coordinates of non-white pixels (actual content)
    non_white_coords = np.where(~white_mask)
    
    # If no content found (completely white image), return original
    if len(non_white_coords[0]) == 0:
        return image
    
    # Find the bounding box of actual content for THIS image
    y_min, y_max = np.min(non_white_coords[0]), np.max(non_white_coords[0])
    x_min, x_max = np.min(non_white_coords[1]), np.max(non_white_coords[1])
    
    # Add small margin
    margin = 5
    y_min = max(0, y_min - margin)
    y_max = min(image.height, y_max + 1 + margin)
    x_min = max(0, x_min - margin)
    x_max = min(image.width, x_max + 1 + margin)
    
    # Crop to the content area
    return image.crop((x_min, y_min, x_max, y_max))

def process_pdf_with_logos(pdf_file, logo_states, white_threshold, removal_method, cropping_method, main_progress, sub_progress, time_tracker):
    """Process all pages with quad logo removal and cropping with detailed progress"""
    processed_images = []
    
    pdf_data = pdf_file.getvalue()
    doc = fitz.open(stream=pdf_data, filetype="pdf")
    total_pages = len(doc)
    
    start_time = time.time()
    
    for page_num in range(total_pages):
        # Update main progress
        main_progress.progress((page_num) / total_pages, text=f"🔄 Processing page {page_num + 1}/{total_pages}")
        
        page = doc[page_num]
        pix = page.get_pixmap()
        img_data = pix.tobytes("ppm")
        pil_image = Image.open(io.BytesIO(img_data))
        
        # Step 1: Logo Removal (all 4 logos)
        sub_progress.progress(0.2, text=f"Removing logos...")
        for i in range(1, 5):
            if logo_states[f'logo{i}_enabled'] and logo_states[f'logo{i}_coords']:
                pil_image = remove_logo_precise(pil_image, logo_states[f'logo{i}_coords'], removal_method)
        
        # Step 2: Cropping
        sub_progress.progress(0.6, text=f"Cropping image...")
        if cropping_method == "vertical":
            pil_image = crop_vertical_only(pil_image, white_threshold)
        elif cropping_method == "horizontal":
            pil_image = crop_horizontal_only(pil_image, white_threshold)
        elif cropping_method == "both":
            pil_image = crop_both_fixed(pil_image, white_threshold)
        # else "none" - no cropping
        
        # Step 3: Finalize
        sub_progress.progress(1.0, text=f"Finalizing page {page_num + 1}...")
        processed_images.append(pil_image)
        
        # Estimate time remaining
        elapsed_time = time.time() - start_time
        pages_processed = page_num + 1
        time_per_page = elapsed_time / pages_processed
        remaining_pages = total_pages - pages_processed
        estimated_remaining = time_per_page * remaining_pages
        
        time_tracker.text(f"⏱️ Estimated time remaining: {estimated_remaining:.1f}s")
        
        # Reset sub-progress for next page
        if page_num < total_pages - 1:
            sub_progress.progress(0.0, text="Ready for next page...")
    
    doc.close()
    return processed_images

def create_word_document_with_options(images, orientation="portrait", margins="normal"):
    """Create Word document with precise image sizing and clean layout"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is not installed. Please run: pip install python-docx")
    
    doc = Document()
    
    # Configure page layout based on orientation
    section = doc.sections[0]
    
    # Set orientation
    if orientation == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)  # A4 landscape
        section.page_height = Cm(21.0)
        image_width = Cm(14.8)
    else:  # portrait
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)  # A4 portrait
        section.page_height = Cm(29.7)
        image_width = Cm(10.0)
    
    # Set margins
    if margins == "narrow":
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
    elif margins == "none":
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)
    else:  # normal
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # Setup two-column layout
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '360')
    cols.set(qn('w:sep'), 'true')
    
    # Add images in two-column layout
    for i, image in enumerate(images):
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
            image.save(tmp_file.name, 'PNG')
            tmp_path = tmp_file.name
        
        try:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(tmp_path, width=image_width)
        finally:
            os.unlink(tmp_path)
    
    return doc

# Initialize session state for 4 logos
for i in range(1, 5):
    if f'logo{i}_enabled' not in st.session_state:
        st.session_state[f'logo{i}_enabled'] = False
    if f'logo{i}_coords' not in st.session_state:
        st.session_state[f'logo{i}_coords'] = None
if 'all_page_images' not in st.session_state:
    st.session_state.all_page_images = None

# Main App Flow
st.sidebar.header("⚙️ PDF Image Processor 1.1")

uploaded_pdf = st.sidebar.file_uploader("1. Upload PDF", type=["pdf"], key="pdf_uploader")

# Word export settings
word_settings = {}
if DOCX_AVAILABLE:
    st.sidebar.subheader("📝 Word Export Settings")
    word_settings['orientation'] = st.sidebar.radio(
        "Page Orientation:",
        ["portrait", "landscape"],
        format_func=lambda x: f"{x.title()} ({'10cm images' if x == 'portrait' else '14.8cm images'})"
    )
    
    word_settings['margins'] = st.sidebar.radio(
        "Page Margins:",
        ["normal", "narrow", "none"],
        format_func=lambda x: {
            "normal": "Normal (1 inch)",
            "narrow": "Narrow (0.5 inch)", 
            "none": "Minimal (0.2 inch)"
        }[x]
    )

if uploaded_pdf:
    st.sidebar.success("✅ PDF uploaded successfully!")
    
    # Extract all pages for logo setup
    if st.session_state.all_page_images is None:
        with st.spinner("Loading PDF pages for logo setup..."):
            st.session_state.all_page_images = get_all_page_images(uploaded_pdf)
    
    # Step 1: Logo Setup
    st.sidebar.subheader("2. Logo Setup")
    setup_logo = st.sidebar.radio("Logo Removal:", ["No Logo Removal", "4-Logo Setup"])
    
    # Step 2: Processing Settings
    st.sidebar.subheader("3. Processing Settings")
    white_threshold = st.sidebar.slider("White Threshold", 200, 254, 245)
    
    removal_method = st.sidebar.selectbox("Logo Removal Method", 
                                        ["white", "smart"],
                                        format_func=lambda x: "White Fill" if x == "white" else "Smart Background Fill")
    
    # Both-direction cropping as default
    cropping_method = st.sidebar.selectbox(
        "Cropping Method:",
        ["both", "vertical", "horizontal", "none"],
        format_func=lambda x: {
            "both": "Both Directions (Default)",
            "vertical": "Vertical Only (top/bottom)", 
            "horizontal": "Horizontal Only (left/right)",
            "none": "No Cropping"
        }[x],
        index=0  # Default to "both"
    )
    
    # Main content area
    if setup_logo == "4-Logo Setup":
        st.header("🎯 Step 2: 4-Logo Visual Setup")
        
        st.info("""
        **Visual Logo Setup:**
        - **🗺️ Grid Overlay**: Toggle grid lines for precise coordinate selection
        - **🎯 Click Buttons**: Set logo areas at common positions
        - **📐 Coordinate Boxes**: Manual precision editing  
        - **👁️ Live Preview**: Real-time visual feedback with grid
        - **⚡ Quick Actions**: Copy, reset, auto-space logos
        """)
        
        # Prepare logo states
        logo_states = {}
        for i in range(1, 5):
            logo_states[f'logo{i}_enabled'] = st.session_state[f'logo{i}_enabled']
            logo_states[f'logo{i}_coords'] = st.session_state[f'logo{i}_coords']
        
        # Get first page for reference
        first_page_img = st.session_state.all_page_images[0]
        
        # Visual logo setup
        visual_logo_selection(first_page_img, logo_states)
    
    # Step 3: Process PDF
    any_logo_enabled = any(st.session_state.get(f'logo{i}_enabled', False) for i in range(1, 5))
    if setup_logo == "No Logo Removal" or any_logo_enabled:
        st.header("🚀 Step 3: Process PDF")
        
        # Show processing summary
        if setup_logo == "4-Logo Setup":
            logo_summary = []
            for i in range(1, 5):
                if st.session_state[f'logo{i}_enabled']:
                    logo_summary.append(f"Logo {i}")
            
            if logo_summary:
                st.info(f"🔧 Will remove: {', '.join(logo_summary)} from all {len(st.session_state.all_page_images)} pages")
        
        # Show cropping info
        st.info(f"🌐 Cropping: **{cropping_method.upper()}** direction{'s' if cropping_method == 'both' else ''}")
        
        if st.button("🔄 Process All Pages", type="primary", key="process_btn"):
            # Create progress containers
            main_progress = st.progress(0, text="🔄 Starting PDF processing...")
            sub_progress = st.progress(0, text="Initializing...")
            time_tracker = st.empty()
            status_text = st.empty()
            
            try:
                # Prepare logo states for processing
                logo_states = {}
                for i in range(1, 5):
                    logo_states[f'logo{i}_enabled'] = st.session_state.get(f'logo{i}_enabled', False)
                    logo_states[f'logo{i}_coords'] = st.session_state.get(f'logo{i}_coords')
                
                processed_images = process_pdf_with_logos(
                    uploaded_pdf, 
                    logo_states,
                    white_threshold, 
                    removal_method, 
                    cropping_method,
                    main_progress,
                    sub_progress,
                    time_tracker
                )
                
                # Final completion
                main_progress.progress(1.0, text="✅ Processing complete!")
                sub_progress.progress(1.0, text="All pages processed successfully")
                time_tracker.empty()
                status_text.success("🎉 All pages processed successfully!")
                
                st.session_state.processed_images = processed_images
                st.session_state.processing_done = True
                
            except Exception as e:
                main_progress.empty()
                sub_progress.empty()
                time_tracker.empty()
                st.error(f"❌ Processing failed: {str(e)}")
    
    # Show results and downloads
    if st.session_state.get('processing_done', False):
        st.header("📊 Results")
        
        # Show before/after comparison
        if any_logo_enabled:
            st.subheader("Before/After Comparison")
            original_first_page = st.session_state.all_page_images[0]
            processed_first_page = st.session_state.processed_images[0]
            
            col1, col2 = st.columns(2)
            with col1:
                st.image(original_first_page, caption="BEFORE - Original", use_column_width=True)
            with col2:
                st.image(processed_first_page, caption="AFTER - Logo(s) removed + Cropped", use_column_width=True)
        
        # Download section
        st.header("📥 Download Results")
        
        if DOCX_AVAILABLE:
            col1, col2, col3 = st.columns(3)
        else:
            col1, col2 = st.columns(2)
        
        with col1:
            # ZIP download
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
                for i, img in enumerate(st.session_state.processed_images):
                    img_bytes = io.BytesIO()
                    img.save(img_bytes, format='PNG')
                    zip_file.writestr(f"page_{i+1:03d}.png", img_bytes.getvalue())
            
            st.download_button(
                label="💾 Download as ZIP (Images)",
                data=zip_buffer.getvalue(),
                file_name="processed_pages.zip",
                mime="application/zip",
                use_container_width=True
            )
        
        with col2:
            # PDF download
            img_bytes_list = []
            for img in st.session_state.processed_images:
                img_bytes = io.BytesIO()
                img.save(img_bytes, format='PNG')
                img_bytes_list.append(img_bytes.getvalue())
            
            pdf_bytes = img2pdf.convert(img_bytes_list)
            
            st.download_button(
                label="📄 Download as PDF",
                data=pdf_bytes,
                file_name="processed_pages.pdf",
                mime="application/pdf",
                use_container_width=True
            )
        
        if DOCX_AVAILABLE:
            with col3:
                # Word Document download
                try:
                    doc = create_word_document_with_options(
                        st.session_state.processed_images,
                        orientation=word_settings.get('orientation', 'portrait'),
                        margins=word_settings.get('margins', 'normal')
                    )
                    
                    doc_buffer = io.BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)
                    
                    st.download_button(
                        label="📝 Download as Word (2-Column)",
                        data=doc_buffer.getvalue(),
                        file_name="processed_pages.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                except Exception as e:
                    st.error(f"Word export failed: {str(e)}")
        
        # Reset button
        if st.button("🔄 Process Another PDF", type="secondary"):
            st.session_state.clear()
            st.rerun()

else:
    # Welcome screen
    st.info("👆 **Get Started:** Upload a PDF file in the sidebar")
    
    st.markdown("""
    ## 🆕 PDF Image Processor 1.1
    
    ### 🎯 **Visual Logo Selection**
    - **🗺️ Grid Overlay**: Toggle grid lines for precise coordinate selection
    - **🎯 Click Buttons**: Set logo areas at common positions (Top-Left, Top-Right, etc.)
    - **📐 Coordinate Boxes**: Manual precision editing  
    - **👁️ Live Preview**: Real-time visual feedback with grid
    - **⚡ Quick Actions**: Copy logos, reset positions, auto-space
    
    ### 🌐 **Both-Direction Cropping (Default)**
    - Automatically crops **top, bottom, left, and right** white spaces
    - Each page analyzed individually for optimal cropping
    """)
    
    if not DOCX_AVAILABLE:
        st.warning("""
        **To enable Word export:**
        ```bash
        python -m pip install python-docx
        ```
        """)

st.markdown("---")
st.caption("PDF Image Processor 1.1 | Visual Logo Selection + Both-Direction Cropping")
